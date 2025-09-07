# -*- coding: utf-8 -*-
"""
Outils Python pour remplacer les scripts externes (PowerShell).
"""
import csv
import os

try:
    from docx import Document
    from docx.shared import RGBColor
except ImportError:
    Document = None


def add_comments_to_docx(original_docx_path: str, comments_csv_path: str, output_docx_path: str):
    """
    Ajoute des commentaires à un document Word à partir d'un fichier CSV.

    Args:
        original_docx_path: Chemin vers le document DOCX de base (copie découpée).
        comments_csv_path: Chemin vers le fichier CSV contenant les commentaires.
        output_docx_path: Chemin où enregistrer le DOCX final avec commentaires.
    """
    if Document is None:
        raise ImportError("La bibliothèque python-docx est requise. Exécutez `pip install python-docx`.")

    # Charger les commentaires depuis le CSV
    comments_to_add = []
    try:
        with open(comments_csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                comments_to_add.append(row)
    except FileNotFoundError:
        # S'il n'y a pas de commentaires, on copie simplement le document et on termine
        import shutil
        shutil.copy(original_docx_path, output_docx_path)
        return

    doc = Document(original_docx_path)

    # Créer un mapping de texte de paragraphe pour une recherche plus rapide
    # On normalise les espaces pour plus de robustesse
    para_texts = [' '.join(p.text.split()) for p in doc.paragraphs]

    for comment_info in comments_to_add:
        anchor = ' '.join(comment_info.get("ancre_textuelle", "").split())
        comment_text = comment_info.get("commentaire", "")
        gravite = comment_info.get("gravite", "P3")

        if not anchor or not comment_text:
            continue

        # Chercher le paragraphe ancre
        found_para = None
        for i, p_text in enumerate(para_texts):
            if anchor in p_text:
                found_para = doc.paragraphs[i]
                break

        if found_para:
            # Ajouter le commentaire au paragraphe trouvé
            author = f"Agent IA ({gravite})"
            initials = gravite
            found_para.add_comment(comment_text, author=author, initials=initials)

    # Sauvegarder le document final
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    doc.save(output_docx_path)
