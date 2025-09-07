from pathlib import Path
from docx import Document


def main():
    # Prefer the exact specified path first
    preferred = Path("input/prédiag_DECOUPE.docx")
    p = preferred
    if not p.exists():
        # Fallback: find any file in input containing 'DECOUPE'
        for cand in Path('input').glob('*'):
            if 'DECOUPE' in cand.name:
                p = cand
                break

    if not p.exists():
        raise FileNotFoundError("Could not find input/prédiag_DECOUPE.docx or any *DECOUPE* file in input/")

    doc = Document(p)
    lines = []
    lines.append(f"SOURCE_FILE: {p}")

    # Extract paragraphs with indices
    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.replace('\t', '    ').strip()
        lines.append(f"PARA[{i:04d}]: {text}")

    # Extract tables content
    for ti, table in enumerate(doc.tables, 1):
        lines.append(f"TABLE[{ti}]: rows={len(table.rows)} cols={len(table.columns)}")
        for ri, row in enumerate(table.rows, 1):
            cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            lines.append(f"T{ti}R{ri:02d}: {' | '.join(cells)}")

    out = Path('_linearized_DECOUPE.txt')
    out.write_text('\n'.join(lines), encoding='utf-8')
    print(f"Extracted to {out}")


if __name__ == "__main__":
    main()

