#!/usr/bin/env python3
import os
import sys
import json
import csv
import shutil
import subprocess
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

try:
    from docx import Document
except Exception:
    Document = None

MODES = [
    ("offre", "Offre"),
    ("diagnostic", "Diagnostic"),
    ("impacts", "Impacts"),
    ("mesures", "Mesures"),
]

ROOT = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(ROOT, "input")
WORK_DIR = os.path.join(ROOT, "work")
OUTPUT_DIR = os.path.join(ROOT, "output")

os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(WORK_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ------------------------- utilitaires -------------------------


def ts_now():
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def norm_style_name(name: str) -> str:
    if not name:
        return ""
    return name.strip().lower()


def detect_head_level(style_name: str):
    n = norm_style_name(style_name)
    # Support FR/EN: "Titre 1/2/3" or "Heading 1/2/3"
    for prefix in ("heading ", "titre "):
        if n.startswith(prefix):
            try:
                return int(n.split(prefix, 1)[1].split()[0])
            except Exception:
                return None
    return None


class Section:
    def __init__(self, title: str, level: int, start_index: int):
        self.title = title.strip()
        self.level = level
        self.start_index = start_index  # index of paragraph in doc.paragraphs
        self.end_index = None  # to be set after pass
        self.number = None  # computed like 1., 2.1, etc.

    def label(self) -> str:
        num = (self.number or "?")
        return f"{num} {self.title}".strip()


def analyze_sections(docx_path: str):
    if Document is None:
        raise RuntimeError("python-docx manquant. Cliquez sur ‘Installer dépendances…’ ou installez: pip install python-docx lxml")
    doc = Document(docx_path)
    secs = []
    for i, p in enumerate(doc.paragraphs):
        lvl = detect_head_level(getattr(p.style, "name", ""))
        if lvl is not None and lvl >= 1:
            txt = p.text.strip()
            if txt:
                secs.append(Section(txt, lvl, i))
    # determine end_index by next section start
    for idx, s in enumerate(secs):
        s.end_index = (secs[idx + 1].start_index if idx + 1 < len(secs) else len(doc.paragraphs))
    # compute numbering by counters
    counters = {}
    for s in secs:
        for k in list(counters.keys()):
            if k > s.level:
                counters.pop(k, None)
        counters[s.level] = counters.get(s.level, 0) + 1
        parts = [str(counters.get(l, 0)) for l in range(1, s.level + 1)]
        s.number = ".".join(parts)
    return secs


def filter_paragraphs_by_sections(docx_path: str, chosen: list, sections: list) -> str:
    """
    Build a new DOCX with only the paragraphs belonging to selected sections.
    chosen: indices into `sections` list.
    Returns path to new docx.
    """
    doc = Document(docx_path)
    new = Document()
    included_ranges = []
    chosen_set = set(chosen)
    # Expand: if a parent is chosen, include all its content until next section
    for i, s in enumerate(sections):
        if i in chosen_set:
            start = s.start_index
            end = s.end_index
            included_ranges.append((start, end))
    if not included_ranges:
        included_ranges = [(0, len(doc.paragraphs))]

    def add_paragraph_like(p_src):
        text = p_src.text
        if not text and not p_src.runs:
            new.add_paragraph("")
            return
        p = new.add_paragraph()
        try:
            p.style = p_src.style
        except Exception:
            pass
        for r in p_src.runs:
            nr = p.add_run(r.text)
            try:
                nr.bold = r.bold
                nr.italic = r.italic
                nr.underline = r.underline
            except Exception:
                pass

    for (a, b) in included_ranges:
        for i in range(a, b):
            add_paragraph_like(doc.paragraphs[i])

    base = os.path.splitext(os.path.basename(docx_path))[0]
    out = os.path.join(WORK_DIR, f"{base}_SECTIONS_{ts_now()}.docx")
    new.save(out)
    return out


def docx_to_markdown(docx_path: str) -> str:
    doc = Document(docx_path)
    lines = []
    for p in doc.paragraphs:
        lvl = detect_head_level(getattr(p.style, "name", ""))
        t = (p.text or "").strip()
        if not t:
            lines.append("")
            continue
        if lvl is not None and 1 <= lvl <= 6:
            lines.append("#" * lvl + " " + t)
        else:
            lines.append(t)
    return "\n".join(lines) + "\n"


def sanitize_text(s: str) -> str:
    import re
    s = s.replace("\u00A0", " ")
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"\s+([,;:!?])", r"\1", s)
    s = re.sub(r"\(\s+", "(", s)
    s = re.sub(r"\s+\)", ")", s)
    return s.strip()


ALLOWED_CATEGORIES = ["coherence", "methodologie", "reglementaire", "carto", "redaction"]


def load_checklist(mode: str):
    path = os.path.join(ROOT, "modes", mode, "instructions", "checklist.md")
    try:
        with open(path, "r", encoding="utf-8") as f:
            lines = [ln.strip(" \t-•") for ln in f.read().splitlines()]
        return [l for l in lines if l and not l.startswith("#")]
    except Exception:
        return []


def classify_comment_from_text(txt: str, mode: str):
    t = txt.lower()
    if any(k in t for k in ["méthod", "methodo", "protocole", "phases", "jalon"]):
        return ("P2", "methodologie")
    if any(k in t for k in ["réglement", "natura 2000", "code", "loi", "eau"]):
        return ("P2", "reglementaire")
    if any(k in t for k in ["carte", "carto", "figure", "légende"]):
        return ("P3", "carto")
    if any(k in t for k in ["cohér", "coherence", "cohérence"]):
        return ("P2", "coherence")
    return ("P3", "redaction")


def generate_review(md_in: str, mode: str):
    """
    Returns revised_md, comments rows.
    comments rows: list of dict with keys ancre_textuelle, commentaire, gravite, categorie
    """
    import re
    lines = md_in.splitlines()
    revised_lines = []
    comments = []
    checklist = load_checklist(mode)

    used_anchors = {}

    def anchor_for(text: str) -> str:
        base = sanitize_text(text)[:20]
        base = re.sub(r"[^A-Za-z0-9À-ÿ\-\s]", "", base)
        base = re.sub(r"\s+", " ", base).strip()
        if not base:
            base = "ancre"
        k = base
        i = 1
        while k.lower() in used_anchors:
            i += 1
            k = f"{base} {i}"
        used_anchors[k.lower()] = True
        return k

    full_text = "\n".join(lines).lower()
    # Find first heading text to anchor generic comments if any
    first_heading_text = None
    for ln_h in lines:
        if ln_h.startswith("#"):
            first_heading_text = sanitize_text(ln_h.lstrip("#").strip())
            break
    for item in checklist:
        key = item.split(":")[0]
        if key and key.lower() not in full_text:
            grav, cat = classify_comment_from_text(item, mode)
            comments.append(
                {
                    "ancre_textuelle": (first_heading_text or "Introduction"),
                    "commentaire": f"Vérifier couverture de la checklist: '{item}'.",
                    "gravite": grav,
                    "categorie": cat,
                }
            )

    for ln in lines:
        raw = ln.rstrip()
        if not raw.strip():
            revised_lines.append("")
            continue
        if raw.startswith("#"):
            parts = raw.split(" ", 1)
            title = sanitize_text(parts[1] if len(parts) > 1 else raw.lstrip("#"))
            revised_lines.append(parts[0] + " " + title)
            continue
        cleaned = sanitize_text(raw)
        revised_lines.append(cleaned)

        txt_low = cleaned.lower()
        make_comment = None
        if len(cleaned) > 600:
            make_comment = ("P3", "redaction", "Paragraphe très long: envisager de le scinder pour la lisibilité.")
        if any(x in txt_low for x in ["tbd", "???", "à définir", "a definir", "xxx"]):
            make_comment = ("P1", "coherence", "Marqueur d'incertain repéré (TBD/??/à définir): préciser ou retirer.")
        if ("carte" in txt_low or "figure" in txt_low) and not re.search(r"\b(\d+)\b", cleaned):
            make_comment = ("P3", "carto", "Référence à une carte/figure sans identifiant: ajouter le numéro.")

        if make_comment:
            grav, cat, note = make_comment
            if cat not in ALLOWED_CATEGORIES:
                cat = "redaction"
            comments.append(
                {
                    "ancre_textuelle": anchor_for(cleaned[:80]),
                    "commentaire": note,
                    "gravite": grav,
                    "categorie": cat,
                }
            )

    revised_md = "\n".join(revised_lines) + "\n"
    return revised_md, comments


def markdown_to_docx(md_text: str, out_path: str):
    doc = Document()
    for raw in md_text.splitlines():
        if not raw.strip():
            doc.add_paragraph("")
            continue
        if raw.startswith("#"):
            level = len(raw) - len(raw.lstrip("#"))
            title = raw[level:].strip()
            p = doc.add_paragraph(title)
            try:
                p.style = f"Heading {min(level,6)}"
            except Exception:
                pass
        elif raw.startswith("- "):
            doc.add_paragraph(raw[2:])
        else:
            doc.add_paragraph(raw)
    doc.save(out_path)


def write_comments_csv(rows, path: str):
    fieldnames = ["ancre_textuelle", "commentaire", "gravite", "categorie"]
    norm_rows = []
    for r in rows:
        rr = {k: str(r.get(k, "")) for k in fieldnames}
        if rr["gravite"] not in ("P1", "P2", "P3"):
            rr["gravite"] = "P3"
        if rr["categorie"] not in ALLOWED_CATEGORIES:
            rr["categorie"] = "redaction"
        if rr["ancre_textuelle"]:
            norm_rows.append(rr)
    with open(path, "w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        for r in norm_rows:
            w.writerow(r)


def run_compare_and_comment(original_docx: str, revised_docx: str, comments_csv: str, output_docx: str):
    """Run the PowerShell compare script and surface detailed errors.

    Uses -NoProfile and captures stdout/stderr so that any failure inside
    PowerShell/Word is clearly shown to the user instead of only the exit code.
    Also normalizes paths to avoid issues with spaces, accents and '&'.
    """
    ps1 = os.path.join(ROOT, "tools", "compare_and_comment.ps1")
    if not os.path.exists(ps1):
        raise RuntimeError("tools/compare_and_comment.ps1 manquant")

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_docx), exist_ok=True)

    # Normalize to absolute paths
    ps1 = os.path.normpath(os.path.abspath(ps1))
    original_docx = os.path.normpath(os.path.abspath(original_docx))
    revised_docx = os.path.normpath(os.path.abspath(revised_docx))
    comments_csv = os.path.normpath(os.path.abspath(comments_csv))
    output_docx = os.path.normpath(os.path.abspath(output_docx))

    # Quick sanity checks for clearer error messages
    for p, label in [
        (original_docx, "OriginalDocx"),
        (revised_docx, "RevisedDocx"),
        (comments_csv, "CommentsCsv"),
    ]:
        if not os.path.exists(p):
            raise RuntimeError(f"Fichier introuvable pour {label}: {p}")

    cmd = [
        "powershell.exe",
        "-NoLogo",
        "-NoProfile",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        ps1,
        "-OriginalDocx",
        original_docx,
        "-RevisedDocx",
        revised_docx,
        "-CommentsCsv",
        comments_csv,
        "-OutputDocx",
        output_docx,
    ]
    try:
        # Capture raw bytes to avoid UnicodeDecodeError with Windows PowerShell streams
        subprocess.run(cmd, check=True, capture_output=True, text=False)
    except subprocess.CalledProcessError as e:
        # Decode PowerShell output robustly
        if isinstance(e.stdout, (bytes, bytearray)):
            try:
                so = (e.stdout or b"").decode("utf-16le")
            except Exception:
                try:
                    so = (e.stdout or b"").decode("utf-8")
                except Exception:
                    so = (e.stdout or b"").decode("cp1252", errors="replace")
        else:
            so = e.stdout or ""
        if isinstance(e.stderr, (bytes, bytearray)):
            try:
                se = (e.stderr or b"").decode("utf-16le")
            except Exception:
                try:
                    se = (e.stderr or b"").decode("utf-8")
                except Exception:
                    se = (e.stderr or b"").decode("cp1252", errors="replace")
        else:
            se = e.stderr or ""
        details = []
        if so.strip():
            details.append("STDOUT:\n" + so.strip())
        if se.strip():
            details.append("STDERR:\n" + se.strip())
        det = ("\n\n".join(details)).strip()
        raise RuntimeError(f"Word Compare a �%chou� (code {e.returncode}).\n\n{det}".rstrip()) from e


# ---------------------- Fenêtre de sélection ----------------------

class SectionsDialog(tk.Toplevel):
    def __init__(self, master, sections, preselected_idx=None):
        super().__init__(master)
        self.title("Sélection des sections")
        self.geometry("560x420")
        self.resizable(True, True)
        self.sections = sections
        self.result = None
        self.selected = set(preselected_idx or [])

        frm = tk.Frame(self)
        frm.pack(fill="both", expand=True, padx=10, pady=10)

        # Listbox multisélection (EXTENDED => Shift pour plage, Ctrl pour ajout)
        self.lb = tk.Listbox(frm, selectmode=tk.EXTENDED, activestyle="none")
        sb = ttk.Scrollbar(frm, orient="vertical", command=self.lb.yview)
        self.lb.configure(yscrollcommand=sb.set)
        self.lb.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        # Remplir
        for i, s in enumerate(sections):
            indent = "    " * (max(0, s.level - 1))
            self.lb.insert(tk.END, f"{indent}{s.label()}")
            if i in self.selected:
                self.lb.selection_set(i)
        # Click sans modificateur => sélectionner le bloc (section + descendants)
        self.lb.bind("<ButtonRelease-1>", self.on_click_block)

        # Actions
        btns = tk.Frame(self)
        btns.pack(fill="x", padx=10, pady=(0, 10))
        tk.Button(btns, text="Tout cocher", command=self.sel_all).pack(side="left")
        tk.Button(btns, text="Tout décocher", command=self.sel_none).pack(side="left", padx=6)
        tk.Button(btns, text="Valider", command=self.on_ok).pack(side="right")
        tk.Button(btns, text="Annuler", command=self.on_cancel).pack(side="right", padx=6)

        self.bind("<Return>", lambda e: self.on_ok())
        self.bind("<Escape>", lambda e: self.on_cancel())
        self.transient(master)
        self.grab_set()
        self.lb.focus_set()

    def sel_all(self):
        self.lb.select_set(0, tk.END)

    def sel_none(self):
        self.lb.select_clear(0, tk.END)

    def on_ok(self):
        self.result = list(self.lb.curselection())
        self.destroy()

    def on_cancel(self):
        self.result = None
        self.destroy()

    def on_click_block(self, event):
        # Respecter Shift/Ctrl (sélection étendue native)
        shift = (event.state & 0x0001) != 0
        ctrl = (event.state & 0x0004) != 0
        if shift or ctrl:
            return
        idx = self.lb.nearest(event.y)
        if idx < 0:
            return
        # Déterminer le bloc [idx .. endDesc]
        cur_level = self.sections[idx].level
        end = idx + 1
        while end < len(self.sections) and self.sections[end].level > cur_level:
            end += 1
        block = range(idx, end)
        # Toggle du bloc (si tout déjà sélectionné -> on désélectionne)
        sel = set(self.lb.curselection())
        if all(i in sel for i in block):
            for i in block:
                self.lb.selection_clear(i)
        else:
            # Remplacer par défaut (comportement simple et prévisible)
            self.lb.selection_clear(0, tk.END)
            for i in block:
                self.lb.selection_set(i)


# ------------------------- GUI -------------------------


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Relecture IA — Assistant graphique")
        self.geometry("1000x760")
        self.resizable(True, True)

        self.source_path = None  # chemin vers le fichier source original
        self.copy_relpath = None  # chemin relatif vers la copie dans input/
        self.output_dir = OUTPUT_DIR  # dossier de sortie choisi
        self.mode = None  # 'offre'|'diagnostic'|'impacts'|'mesures'

        self.sections = []  # list[Section]
        self.section_vars = []  # list[tk.BooleanVar]

        self._build_ui()

    # --- deps ---
    def ensure_docx(self) -> bool:
        global Document
        if Document is not None:
            return True
        try:
            import importlib
            Document = importlib.import_module('docx').Document
            return True
        except Exception:
            pass
        # Try on-the-fly install
        try:
            self.log("Installation de python-docx (et lxml)…")
            subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "lxml"], check=True)
            import importlib
            Document = importlib.import_module('docx').Document
            self.log("python-docx installé avec succès.")
            return True
        except Exception as e:
            self.log(f"Échec installation python-docx: {e}")
            return False

    def log(self, msg: str):
        self.log_txt.configure(state="normal")
        self.log_txt.insert(tk.END, f"{datetime.now().strftime('%H:%M:%S')} — {msg}\n")
        self.log_txt.see(tk.END)
        self.log_txt.configure(state="disabled")
        self.update_idletasks()

    def _build_ui(self):
        pad = {"padx": 10, "pady": 8}

        # 1) Fichier source
        frame1 = tk.LabelFrame(self, text="1) Fichier source")
        frame1.pack(fill="x", **pad)
        self.src_var = tk.StringVar(value="Aucun fichier sélectionné")
        tk.Label(frame1, textvariable=self.src_var, anchor="w").pack(fill="x", padx=10, pady=5)
        btns1 = tk.Frame(frame1)
        btns1.pack(anchor="w", padx=10, pady=5)
        tk.Button(btns1, text="Sélectionner le fichier Word…", command=self.choose_source).pack(side="left", padx=5)
        tk.Button(btns1, text="Ouvrir le dossier", command=self.open_source_folder).pack(side="left", padx=5)
        tk.Label(
            frame1,
            text="Le fichier original ne sera jamais modifié. Le traitement s’effectue sur une copie.",
            fg="#555",
        ).pack(anchor="w", padx=10, pady=5)

        # 2) Sections du document
        frame2 = tk.LabelFrame(self, text="2) Sections du document")
        frame2.pack(fill="both", expand=False, **pad)
        btns2 = tk.Frame(frame2)
        btns2.pack(anchor="w", padx=10, pady=5)
        tk.Button(btns2, text="Afficher les sections", command=self.show_sections).pack(side="left", padx=5)
        tk.Button(btns2, text="Installer dépendances…", command=self.install_deps).pack(side="left", padx=5)
        self.sections_count_var = tk.StringVar(value="0 section sélectionnée")
        tk.Label(btns2, textvariable=self.sections_count_var).pack(side="left", padx=10)

        # Zone placeholder (reste vide; la sélection se fait dans une fenêtre dédiée scrollable)
        self.sections_canvas = tk.Canvas(frame2, height=120)
        self.sections_canvas.pack(fill="x", padx=10, pady=5)
        self.sections_frame = tk.Frame(self.sections_canvas)
        self.sections_canvas.create_window((0, 0), window=self.sections_frame, anchor="nw")

        # 3) Mode de relecture
        frame3 = tk.LabelFrame(self, text="3) Mode de relecture")
        frame3.pack(fill="x", **pad)
        self.mode_buttons = {}
        btnrow = tk.Frame(frame3)
        btnrow.pack(anchor="w", padx=10, pady=5)
        for key, label in MODES:
            b = tk.Button(btnrow, text=label, width=20, command=lambda k=key: self.set_mode(k))
            b.pack(side="left", padx=6, pady=3)
            self.mode_buttons[key] = b

        # 4) Dossier de sortie
        frame4 = tk.LabelFrame(self, text="4) Dossier de sortie")
        frame4.pack(fill="x", **pad)
        self.out_var = tk.StringVar(value=self.output_dir)
        tk.Label(frame4, textvariable=self.out_var, anchor="w").pack(fill="x", padx=10, pady=5)
        btns4 = tk.Frame(frame4)
        btns4.pack(anchor="w", padx=10, pady=5)
        tk.Button(btns4, text="Choisir le dossier de sortie…", command=self.choose_output_dir).pack(side="left", padx=5)
        tk.Button(btns4, text="Ouvrir le dossier de sortie", command=self.open_output_dir).pack(side="left", padx=5)

        # 5) Lancer l’analyse
        frame5 = tk.LabelFrame(self, text="5) Lancer l’analyse")
        frame5.pack(fill="both", expand=True, **pad)
        tk.Button(frame5, text="Lancer l’analyse", command=self.launch_analysis).pack(anchor="w", padx=10, pady=5)

        tk.Label(frame5, text="Journal d’exécution").pack(anchor="w", padx=10)
        self.log_txt = tk.Text(frame5, height=10, wrap="word", state="disabled")
        self.log_txt.pack(fill="both", expand=True, padx=10, pady=5)

        tk.Label(frame5, text="Prompt opérationnel (audit)").pack(anchor="w", padx=10)
        self.prompt_txt = tk.Text(frame5, height=10, wrap="word")
        self.prompt_txt.configure(state="disabled")
        self.prompt_txt.pack(fill="both", expand=True, padx=10, pady=5)

    # Actions
    def choose_source(self):
        path = filedialog.askopenfilename(filetypes=[("Documents Word", "*.docx")])
        if not path:
            return
        self.source_path = os.path.abspath(path)
        self.src_var.set(self.source_path)
        try:
            self.copy_relpath = self._copy_to_input(self.source_path)
            self.log(f"Copie créée: {self.copy_relpath}")
        except Exception as e:
            messagebox.showerror("Erreur de copie", str(e))
            self.copy_relpath = None

    def open_source_folder(self):
        path = self.source_path
        if not path:
            messagebox.showinfo("Info", "Aucun fichier sélectionné.")
            return
        folder = os.path.dirname(path)
        try:
            os.startfile(folder)
        except Exception as e:
            messagebox.showerror("Erreur", str(e))

    def set_mode(self, key):
        self.mode = key
        for k, b in self.mode_buttons.items():
            if k == key:
                b.configure(bg="#2e86de", fg="white", activebackground="#1e5aa6")
            else:
                b.configure(bg=self.cget("bg"), fg="black", activebackground=self.cget("bg"))

    def choose_output_dir(self):
        d = filedialog.askdirectory()
        if not d:
            return
        self.output_dir = os.path.abspath(d)
        os.makedirs(self.output_dir, exist_ok=True)
        self.out_var.set(self.output_dir)

    def open_output_dir(self):
        d = self.output_dir or OUTPUT_DIR
        os.makedirs(d, exist_ok=True)
        try:
            os.startfile(d)
        except Exception as e:
            messagebox.showerror("Erreur", str(e))

    def _copy_to_input(self, src_path):
        os.makedirs(INPUT_DIR, exist_ok=True)
        base = os.path.basename(src_path)
        name, ext = os.path.splitext(base)
        candidate = os.path.join(INPUT_DIR, f"{name}_copie_{ts_now()}{ext}")
        shutil.copy2(src_path, candidate)
        rel = os.path.relpath(candidate, ROOT).replace("\\", "/")
        return rel

    def clear_sections_ui(self):
        for w in list(self.sections_frame.children.values()):
            w.destroy()
        self.section_vars.clear()
        self.sections_count_var.set("0 section sélectionnée")

    def show_sections(self):
        if not self.copy_relpath:
            messagebox.showwarning("Manque fichier", "Sélectionnez d’abord un fichier Word.")
            return
        if not self.ensure_docx():
            messagebox.showerror("Dépendances manquantes", "python-docx introuvable. Essayez ‘Installer dépendances…’ ou exécutez:\npython -m pip install python-docx lxml")
            return
        self.clear_sections_ui()
        try:
            abs_path = os.path.join(ROOT, self.copy_relpath)
            secs = analyze_sections(abs_path)
            if not secs:
                self.log("Aucune table des matières détectée. Bascule vers les styles de titres.")
                secs = analyze_sections(abs_path)
            self.sections = secs
            # Ouvrir la fenêtre modale scrollable pour sélectionner
            pre = [i for i, v in enumerate(self.section_vars) if v.get()] if self.section_vars else []
            dlg = SectionsDialog(self, secs, pre)
            self.wait_window(dlg)
            if dlg.result is not None:
                # Recréer les vars selon résultat
                self.section_vars = []
                selected = set(dlg.result)
                for i in range(len(secs)):
                    self.section_vars.append(tk.BooleanVar(value=(i in selected)))
                # Afficher un petit résumé (les 6 premiers labels) dans la zone placeholder
                for w in list(self.sections_frame.children.values()):
                    w.destroy()
                summary = [secs[i].label() for i in sorted(selected)][:6]
                if summary:
                    for lab in summary:
                        tk.Label(self.sections_frame, text=f"☑ {lab}", anchor="w").pack(anchor="w")
                    if len(selected) > 6:
                        tk.Label(self.sections_frame, text="…", anchor="w").pack(anchor="w")
                else:
                    tk.Label(self.sections_frame, text="(document entier)", anchor="w").pack(anchor="w")
            self.update_sections_count()
        except Exception as e:
            messagebox.showerror("Erreur", f"Impossible d’analyser les sections:\n{e}")

    def install_deps(self):
        ok = self.ensure_docx()
        if ok:
            messagebox.showinfo("OK", "Dépendances installées.")
        else:
            messagebox.showerror("Erreur", "Impossible d’installer python-docx automatiquement. Installez manuellement:\npython -m pip install python-docx lxml")

    def update_sections_count(self):
        n = sum(1 for v in self.section_vars if v.get())
        self.sections_count_var.set(f"{n} section(s) sélectionnée(s)")

    def build_prompt_preview(self, rel_docx: str, selected_labels: list, mode: str) -> str:
        mode_label = dict(MODES)[mode]
        checklist = f"modes/{mode}/instructions/checklist.md"
        refs_dir = f"modes/{mode}/refs"
        lines = []
        lines.append(f"Objectif: Relecture ‘{mode_label}’ du document {rel_docx}.")
        if selected_labels:
            lines.append("Sections: " + ", ".join(selected_labels))
        else:
            lines.append("Sections: document entier")
        lines.append(f"Checklist: {checklist}")
        lines.append(f"Références locales: {refs_dir}")
        return "\n".join(lines)

    def launch_analysis(self):
        if not self.copy_relpath:
            messagebox.showwarning("Manque fichier", "Veuillez sélectionner un fichier Word.")
            return
        if not self.mode:
            messagebox.showwarning("Manque mode", "Veuillez sélectionner un mode de relecture.")
            return
        if not self.ensure_docx():
            messagebox.showerror("Dépendances manquantes", "python-docx introuvable. Essayez ‘Installer dépendances…’.")
            return
        selected_idx = [i for i, v in enumerate(self.section_vars) if v.get()]
        if not selected_idx:
            if not messagebox.askyesno(
                "Document entier", "Aucune section cochée. Voulez-vous traiter le document entier ?"
            ):
                return

        # A — Préparation
        self.log("Préparation…")
        ts = ts_now()
        os.makedirs(WORK_DIR, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)

        abs_copy = os.path.join(ROOT, self.copy_relpath)
        base_name = os.path.splitext(os.path.basename(abs_copy))[0]
        work_copy = os.path.join(WORK_DIR, f"{base_name}_copie_{ts}.docx")
        shutil.copy2(abs_copy, work_copy)

        session = {
            "timestamp": datetime.now().isoformat(),
            "source_docx": self.copy_relpath,
            "mode": self.mode,
            "sections": [self.sections[i].label() for i in selected_idx],
            "output_dir": self.output_dir,
        }
        with open(os.path.join(WORK_DIR, "session.json"), "w", encoding="utf-8") as f:
            json.dump(session, f, ensure_ascii=False, indent=2)

        prompt = self.build_prompt_preview(self.copy_relpath, session["sections"], self.mode)
        self.prompt_txt.configure(state="normal")
        self.prompt_txt.delete("1.0", tk.END)
        self.prompt_txt.insert(tk.END, prompt)
        self.prompt_txt.configure(state="disabled")

        # B — Découpage par sections
        try:
            self.log("Découpage par sections…")
            cut_docx = filter_paragraphs_by_sections(work_copy, selected_idx, self.sections)
        except Exception as e:
            messagebox.showerror("Erreur — Découpage", str(e))
            return

        # C — Relecture locale
        try:
            self.log("Conversion DOCX -> Markdown (travail)…")
            md_src = docx_to_markdown(cut_docx)
            self.log("Application de la checklist et génération des commentaires…")
            revised_md, comments = generate_review(md_src, self.mode)
            md_path = os.path.join(WORK_DIR, "rapport_revise.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(revised_md)
            csv_path = os.path.join(WORK_DIR, "commentaires.csv")
            write_comments_csv(comments, csv_path)
        except Exception as e:
            messagebox.showerror("Erreur — Relecture", str(e))
            return

        # D — Livrable final Word
        try:
            self.log("Conversion Markdown -> DOCX révisé…")
            revised_docx = os.path.join(WORK_DIR, "rapport_revise.docx")
            markdown_to_docx(revised_md, revised_docx)

            self.log("Comparaison et insertion des commentaires…")
            # Avoid very long filenames which can exceed MAX_PATH in some setups
            base_short = (base_name[:60]).rstrip(" _-.")
            out_name = f"{base_short}_AI_suivi+commentaires_{ts_now()}.docx"
            out_docx = os.path.join(self.output_dir, out_name)
            run_compare_and_comment(cut_docx, revised_docx, csv_path, out_docx)
        except subprocess.CalledProcessError as pe:
            messagebox.showerror("Erreur — Word Compare", f"Échec de la comparaison Word: {pe}")
            return
        except Exception as e:
            messagebox.showerror("Erreur — Livrable", str(e))
            return

        self.log("Livrable prêt.")
        self.log(f"Fichier final: {out_docx}")
        messagebox.showinfo("Terminé", f"Livrable généré:\n{out_docx}")


if __name__ == "__main__":
    App().mainloop()
