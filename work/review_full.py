# -*- coding: utf-8 -*-
import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

SRC = 'input/prédiag_DECOUPE.docx'
OUT = r"c:\Users\utilisateur\Mon Drive\1 - Bota & Travail\+++++++++  BOTA  +++++++++\---------------------- 3) BDD\PYTHON\0) Scripts Python\Agent-relecture\Agent-relecture\output\prédiag_AI_suivi+commentaires.docx"

doc = Document(SRC)

# Utils

def add_comment_after_par(p, text):
    # Fallback: append a new paragraph at end (since python-docx lacks comments API)
    note = doc.add_paragraph('[Commentaire] ' + text)
    r = note.runs[0]
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return note

# Seal scope
head_p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph('')
add_comment_after_par(head_p, "Revue DIAGNOSTIC appliquée uniquement à input/prédiag_DECOUPE.docx. Source confirmée.")

# Checklist anchors
checks = [
    ('Périmètre', "Définir limites spatiales/temporaires, emprises, aires d\'étude."),
    ('Méthodologie', "Sources, protocoles, périodes de prospection, limites de validité."),
    ('Données', "Tracer sources, millésimes, dispositifs de collecte."),
    ('Sensibilités', "Habitats, espèces protégées, ZNIEFF, Natura 2000, continuités."),
    ('Impacts', "Par phase et scénario, intensité/durée/réversibilité."),
    ('Mesures', "ERC: Eviter, Réduire, Compenser; hiérarchisation."),
    ('Suivi', "Indicateurs, modalités de contrôle et calendrier."),
]

full_text = "\n".join(p.text for p in doc.paragraphs)
lt = full_text.lower()
for key, guidance in checks:
    if key.lower() not in lt:
        add_comment_after_par(doc.paragraphs[-1], f"Rubrique non repérée: {key}. Attendu: {guidance}")

# Heuristics flags
patterns = [
    (r"\bpeut(?: |-)être\b", "Formulation incertaine. Préciser hypothèse et source."),
    (r"\bprobable(?:ment)?\b", "Qualificatif vague. Justifier ou nuancer."),
    (r"\bimpact(s)? faible(s)?\b", "Qualifier par intensité/durée/étendue + source."),
    (r"\baucun impact\b", "Forte assertion: fournir démonstration et références."),
    (r"\bselon nos connaissances\b", "Préciser période et corpus de référence."),
]

for p in doc.paragraphs:
    t = p.text or ''
    for rx, note in patterns:
        if re.search(rx, t, flags=re.I):
            add_comment_after_par(p, note)
            break

# Structure hints if nothing obvious
sections = [
    'Périmètre de l\'étude',
    'Méthodologie et données',
    'Sensibilités écologiques et réglementaires',
    'Analyse des impacts',
    'Mesures ERC',
    'Suivi et indicateurs',
]

found_any = any(any(s.lower() in (pp.text or '').lower() for s in sections) for pp in doc.paragraphs)
if not found_any:
    sep = doc.add_paragraph('--- Repères de structuration à intégrer selon checklist ---')
    sep.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    for s in sections:
        doc.add_paragraph(f"A insérer: {s}")

Path(OUT).parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print('OK:', OUT)
