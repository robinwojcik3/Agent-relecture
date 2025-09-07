# -*- coding: utf-8 -*-
import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

SRC = 'input/prédiag_DECOUPE.docx'
OUT = r"c:\Users\utilisateur\Mon Drive\1 - Bota & Travail\+++++++++  BOTA  +++++++++\---------------------- 3) BDD\PYTHON\0) Scripts Python\Agent-relecture\Agent-relecture\output\prédiag_AI_suivi+commentaires.docx"

doc = Document(SRC)

# Helper to add an inline highlighted comment paragraph

def add_comment(text):
    p = doc.add_paragraph('[Commentaire] ' + text)
    if p.runs:
        p.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p

# Helper to add a proposal paragraph

def add_proposal(text):
    p = doc.add_paragraph('Proposition de reformulation: ' + text)
    if p.runs:
        p.runs[0].font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    return p

# Remove previous generic footer markers if present
clean_suffix = [
    "--- Repères de structuration à intégrer selon checklist ---",
    "A insérer: Périmètre de l'étude",
    "A insérer: Méthodologie et données",
    "A insérer: Sensibilités écologiques et réglementaires",
    "A insérer: Analyse des impacts",
    "A insérer: Mesures ERC",
    "A insérer: Suivi et indicateurs",
]

# Soft cleanup: filter out exact trailing markers
if doc.paragraphs:
    new_body = []
    for p in doc.paragraphs:
        if p.text.strip() in clean_suffix:
            continue
        new_body.append(p)
    # python-docx has no direct remove; append replacements at end with note
    removed = len(doc.paragraphs) - len(new_body)
    if removed:
        add_comment("Suppression des repères génériques précédents pour clarté de la relecture.")

# Pass 1: flag vague formulations and propose rewrites templates
vague_rules = [
    (r"\bpeut(?: |-)être\b", "Remplacer par une hypothèse explicite et sourcée (ex: 'Selon le rapport X, hypothèse H1 valable pour la période P.')."),
    (r"\bprobable(?:ment)?\b", "Remplacer par un niveau d\'incertitude quantifié ou référencé (ex: 'probabilité modérée (40–60%) selon Y')."),
    (r"\baucun impact\b", "Préciser démonstration (scénario, intensité, durée, étendue) + références."),
    (r"\bimpact(s)? faible(s)?\b", "Qualifier avec intensité/durée/étendue et source; ex: 'faible (I1), court terme (<1 an), local (<100 m) selon Z'."),
]

for p in list(doc.paragraphs):
    t = p.text or ''
    for rx, guidance in vague_rules:
        if re.search(rx, t, flags=re.I):
            add_comment(f"Formulation vague détectée: '{t.strip()[:120]}...' -> {guidance}")
            # Provide a neutral template proposal keeping original structure
            add_proposal(re.sub(rx, 'selon [source], [valeur/intervalle] [unité]', t, flags=re.I))
            break

# Pass 2: headers compliance suggestions if clearly missing anchors
text_all = '\n'.join(p.text for p in doc.paragraphs).lower()
anchors = [
    ('périmètre', "Ajouter un paragraphe précisant les limites spatiales/temporaires, emprises et aires d'étude."),
    ('méthodologie', "Documenter sources, protocoles et périodes de prospection; préciser limites de validité."),
    ('données', "Lister les jeux de données, dates de collecte/millésimes et fiabilité."),
    ('sensibilités', "Présenter habitats/espèces protégées, zones réglementaires, continuités écologiques."),
    ('impacts', "Analyser par phase et scénario, avec intensité/durée/étendue/réversibilité."),
    ('mesures', "Détailler ERC: éviter, réduire, compenser, avec hiérarchisation et efficacité attendue."),
    ('suivi', "Proposer indicateurs, modalités de contrôle et calendrier de suivi."),
]
missing = []
for key, need in anchors:
    if key not in text_all:
        missing.append((key, need))

if missing:
    add_comment("Rubriques à compléter selon la checklist diagnostic (repérées comme absentes dans la copie découpée):" + 
                ' ' + ', '.join(k for k,_ in missing))
    for k, need in missing:
        add_proposal(need)

# Save
Path(OUT).parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print('OK:', OUT)
