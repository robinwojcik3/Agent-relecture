# -*- coding: utf-8 -*-
import re, os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
SRC = 'input/prédiag_DECOUPE.docx'
OUT = r"c:\Users\utilisateur\Mon Drive\1 - Bota & Travail\+++++++++  BOTA  +++++++++\---------------------- 3) BDD\PYTHON\0) Scripts Python\Agent-relecture\Agent-relecture\output\prédiag_AI_suivi+commentaires.docx"

d = Document(SRC)

def comment(p, txt):
    r = p.add_run(f" [Commentaire: {txt}] ")
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW

checks = [
    ('Périmètre', "Vérifier que les limites d’étude sont explicites."),
    ('Méthodologie', 'Préciser sources, protocoles et périodes de prospection.'),
    ('Sensibilités', 'Identifier habitats, espèces protégées et enjeux.'),
    ('Impacts', 'Qualifier par scénario et intensité.'),
    ('Mesures', 'Eviter, réduire, compenser; hiérarchiser.'),
    ('Données', 'Citer sources et dates des données.'),
]

full_text = "\n".join(p.text for p in d.paragraphs)
for key, note in checks:
    if key.lower() not in full_text.lower():
        p = d.add_paragraph('')
        comment(p, f"Rubrique manquante ou non repérée: {key}. {note}")

comment(d.paragraphs[0] if d.paragraphs else d.add_paragraph(''),
        'Audit automatique: passage initial de vérification des rubriques clés (diagnostic).')

d.save(OUT)
print('ok')
